from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(r"C:\Users\admin\Documents\视频小工具\剪影工坊-产品与开发需求规格.docx")
BLUE = "6C5CE7"
BLUE_DARK = "4536A8"
INK = "20232D"
MUTED = "687083"
LIGHT = "F2F1FF"
GREEN = "27AE7A"
RED = "C94957"
GRAY = "F3F5F8"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(.82)
sec.bottom_margin = Inches(.78)
sec.left_margin = Inches(.86)
sec.right_margin = Inches(.86)
sec.header_distance = Inches(.35)
sec.footer_distance = Inches(.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Microsoft YaHei'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.22

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE_DARK, 16, 7),
    ('Heading 2', 13, BLUE, 11, 5),
    ('Heading 3', 11.5, BLUE_DARK, 8, 4),
]:
    st = styles[name]
    st.font.name = 'Microsoft YaHei'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ['List Bullet', 'List Number']:
    st = styles[name]
    st.font.name = 'Microsoft YaHei'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(.38)
    st.paragraph_format.first_line_indent = Inches(-.18)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.line_spacing = 1.18

header = sec.header.paragraphs[0]
header.text = '剪影工坊  ·  产品与开发需求规格'
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_run = header.runs[0]
header_run.font.name = 'Microsoft YaHei'
header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
header_run.font.size = Pt(8.5)
header_run.font.color.rgb = RGBColor.from_string(MUTED)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('内部开发基准  ·  第 ')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(MUTED)
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
footer._p.append(fld)
footer.add_run(' 页')

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = 'Table Grid'
    if widths is None: widths = [6.5/len(headers)]*len(headers)
    for i, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = Inches(widths[i]); set_cell_shading(cell, BLUE_DARK); set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text); run.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(9.5)
        run.font.name='Microsoft YaHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (cell, text) in enumerate(zip(cells, row)):
            cell.width = Inches(widths[i]); set_cell_margins(cell)
            if ridx % 2: set_cell_shading(cell, 'F8F9FB')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12
            run=p.add_run(str(text)); run.font.size=Pt(9.2); run.font.name='Microsoft YaHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def add_bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet')
    if level: p.paragraph_format.left_indent = Inches(.38 + .28*level)
    p.add_run(text)
    return p

def add_num(text):
    p=doc.add_paragraph(style='List Number'); p.add_run(text); return p

def add_callout(title, text, color=BLUE):
    table=doc.add_table(rows=1, cols=1); table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=table.cell(0,0); cell.width=Inches(6.5); set_cell_shading(cell, LIGHT); set_cell_margins(cell,130,160,130,160)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2)
    a=p.add_run(title+'  '); a.bold=True; a.font.color.rgb=RGBColor.from_string(color)
    b=p.add_run(text); b.font.color.rgb=RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

# Title block
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
r=p.add_run('剪影工坊'); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(27); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE_DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14)
r=p.add_run('产品与开发需求规格'); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(MUTED)
add_table(['文档属性','内容'],[
    ('产品定位','日常视频与音频素材快速处理工具'),
    ('目标平台','Windows 与 macOS 桌面端'),
    ('界面方向','简洁、一屏完成、深色蓝紫主题'),
    ('需求状态','汇总用户已明确提出及最终确认的要求'),
    ('文档用途','开发实现、功能检查与交付验收基准'),
],[1.45,5.05])
add_callout('核心产品语句', '拖入素材，剪一段、排一下、转格式，马上导出。')

doc.add_heading('1. 产品定位与设计边界', level=1)
doc.add_paragraph('剪影工坊不是专业剪辑软件，而是用于临时、快速处理日常素材的轻量桌面工具。功能可以覆盖高频需求，但操作流程必须直接、低学习成本，不能因为功能增加而演变成复杂的多轨剪辑系统。')
doc.add_heading('1.1 设计原则', level=2)
for x in [
    '拖入即编辑，不强制创建工程、素材库或复杂项目结构。',
    '常用功能在一屏内完成；当前步骤无关的设置应隐藏或折叠。',
    '默认参数必须合理，多数用户无需理解专业术语即可直接导出。',
    '实际功能优先于表面展示：不得使用模拟缩略图、示意波形或无后端逻辑的按钮。',
    'UI延续早期版本的简洁、现代、深色风格，不采用专业剪辑软件式复杂布局。',
    '遇到技术问题必须定位根因，必要时查阅官方资料，不以占位方案伪装完成。',
]: add_bullet(x)
doc.add_heading('1.2 明确不做', level=2)
for x in ['多轨视频/音频编辑','关键帧动画','专业调色与特效系统','字幕轨与复杂字幕编辑','素材库和媒体资产管理','复杂混音台','图层、节点或合成系统','需要专业剪辑知识的工作流']:
    add_bullet(x)

doc.add_heading('2. 核心用户流程', level=1)
for x in ['直接打开或拖入一个或多个视频/音频文件。','在预览区拖动入点、出点，准确选择需要的范围。','分割片段、删除片段并用鼠标拖动调整顺序。','按需要调整当前片段的画面或音量。','选择画幅、分辨率、帧率、质量和保存方式。','导出视频或分离/导出音频。']:
    add_num(x)
add_callout('目标', '从导入到导出应尽量控制在四个核心动作内：拖入、裁切、排序、导出。')

doc.add_heading('3. 平台、安装与媒体兼容', level=1)
doc.add_heading('3.1 桌面平台', level=2)
for x in ['Windows提供可安装程序和可直接运行的验证版本。','macOS提供DMG或等效安装包；正式分发时支持签名与公证流程。','FFmpeg随应用打包，普通用户无需单独安装或配置命令行工具。','隐藏Windows原生风格标题栏，使用与主体一致的窗口控制按钮。']:
    add_bullet(x)
doc.add_heading('3.2 视频与相机素材', level=2)
add_table(['类别','要求'],[
    ('常见容器','MP4、MOV、MKV、AVI、WebM、M4V、TS、MPEG/MPG、VOB、3GP'),
    ('索尼相机','XAVC S / XAVC HS（MP4）、XAVC-I（MXF）、AVCHD（MTS/M2TS/M2T）'),
    ('常见编码','H.264、H.265/HEVC、ProRes、DNxHD/DNxHR、MPEG-2等'),
    ('预览代理','无法由Electron直接播放的编码自动生成轻量H.264代理；导出仍读取原文件'),
],[1.35,5.15])
doc.add_heading('3.3 音频素材', level=2)
for x in ['支持MP3、WAV、M4A、AAC、FLAC、OGG等常见音频格式。','允许建立纯音频处理时间轴，裁切与拖动操作和视频保持一致。','视频与纯音频可分别处理；不以复杂多轨混合为目标。']:
    add_bullet(x)

doc.add_heading('4. 导入与预览', level=1)
for x in [
    '空白区域支持点击选择文件，也支持直接拖入文件；一次可拖入多段素材。',
    '时间轴默认“添加”区域本身也必须支持文件拖入。',
    '导入过程中明确显示“分析素材、提取首帧、生成代理、计算波形”等真实状态。',
    '选中任意片段后必须立即显示对应视频或音频预览，不能只验证窗口是否启动。',
    '切换片段、分割、排序、删除或切换时间轴后，不得引用已不存在的片段。',
]: add_bullet(x)
doc.add_heading('4.1 视频缩略图', level=2)
for x in ['每个视频片段必须使用FFmpeg从真实原素材或代理中提取首帧JPEG。','缩略图失败必须显示明确错误或重试状态，不得静默显示永久占位。','分割生成的片段应显示其对应起始位置的帧，或至少可靠显示源素材首帧。','片段名称位于缩略图上方，不放在缩略图侧面。']:
    add_bullet(x)

doc.add_heading('5. 精确裁切与播放', level=1)
doc.add_heading('5.1 入点与出点', level=2)
for x in [
    '只保留入点和出点作为区间控制，不在下方进度条显示额外播放头。',
    '入点与出点在进度条上清晰显示，选中区间使用绿色区分。',
    '入点、出点手柄可用鼠标连续拖动，并显示准确时间。',
    '拖动入点或出点时，上方视频画面必须实时跳转到对应位置。',
    '音频模式中，拖动入出点时长光标必须同步移动到准确位置。',
    '精确时间数字输入移动到入点、出点、分割所在的工具栏右侧；取消时间轴下方独立“精确时间”区域。',
]: add_bullet(x)
doc.add_heading('5.2 播放规则', level=2)
for x in ['播放按钮只播放当前入点至出点的区间。','播放到出点后立即暂停，不继续播放区间外内容。','重新播放时，如果当前位置不在有效区间内，从入点开始。','进度条支持点击和连续拖动定位；拖动时预览实时跟随并保持暂停。','播放按钮颜色使用主体蓝紫色。']:
    add_bullet(x)
doc.add_heading('5.3 分割', level=2)
for x in ['支持在当前定位位置将片段分割为两段。','支持设定入点、出点后按区间裁切片段。','分割后的每个片段都必须能够独立选中、预览、排序、删除和导出。']:
    add_bullet(x)

doc.add_heading('6. 片段时间轴', level=1)
for x in [
    '时间轴为单层片段序列，不引入专业多轨结构。',
    '可拖动片段调整顺序；拖动经过目标位置时，其他片段平滑让位，视觉结果必须直观。',
    '视频片段显示真实首帧；音频片段显示真实波形。',
    '片段宽度应大致体现时长，同时保持日常操作可读性。',
    '片段卡片提供删除入口；时间轴右上角提供关闭/清空当前时间轴选项。',
    '空白区域提供“新建时间轴”按钮，可增加独立空白时间轴。',
    '多时间轴使用轻量标签切换，不纵向堆叠为复杂轨道。',
    '时间轴横向滚动条必须使用与主题一致的蓝紫样式，不显示浏览器默认风格。',
]: add_bullet(x)
doc.add_heading('6.1 多时间轴导出', level=2)
for x in ['存在多条非空时间轴时，每条时间轴分别生成一个输出文件。','文件名应包含时间轴名称并避免覆盖。','空时间轴不生成文件。','导出界面应明确当前将输出的时间轴数量。']:
    add_bullet(x)

doc.add_heading('7. 视频片段调整', level=1)
for x in [
    '画面缩放、旋转90度、水平翻转、垂直翻转只作用于当前选中片段。',
    '调整当前片段时，上方视频预览实时反映变化，其他片段不受影响。',
    '切换片段后，右侧控件显示该片段自己的参数。',
    '导出时分别应用每个片段的独立变换参数。',
]: add_bullet(x)
doc.add_heading('7.1 画幅比例', level=2)
for x in ['支持原始比例、16:9、9:16、1:1。','选择非原始比例后提供“自动裁剪铺满”和“完整画面加黑边”两种方式。','预览画面必须与最终导出的裁剪/黑边策略保持一致。']:
    add_bullet(x)

doc.add_heading('8. 音频处理', level=1)
doc.add_heading('8.1 真实波形', level=2)
for x in [
    '波形必须由FFmpeg解码后的真实PCM采样计算，不使用随机、固定或示意波形。',
    '上方音频编辑窗口显示当前片段完整的详细波形和贯穿区域的长光标。',
    '音频时间轴片段同样显示真实波形的降采样缩略视图。',
    '波形分析失败时显示错误和重试状态，不显示伪造结果。',
]: add_bullet(x)
doc.add_heading('8.2 波形导航', level=2)
for x in ['鼠标滚轮缩放波形。','鼠标中键拖动平移波形视图。','拖动长光标可定位音频并实时更新时间。','底部进度区采用蓝色主题UI，不使用浏览器默认滚动条。']:
    add_bullet(x)
doc.add_heading('8.3 音量与响度', level=2)
for x in [
    '单个音频片段提供音量增益，界面单位使用dB，默认0 dB；建议范围-24 dB至+12 dB，并提供静音。',
    '提供全部音频的统一总音量调节。',
    '提供统一响度开关，使用FFmpeg loudnorm和EBU R128/ITU-R BS.1770测量算法。',
    '网络视频/播客默认目标可设为-16 LUFS，真峰值限制建议-1.5 dBTP。',
    '界面文案应表述为“使用EBU R128测量算法，目标-16 LUFS”，不能误称EBU标准固定要求-16 LUFS。',
]: add_bullet(x)
doc.add_heading('8.4 音频分离', level=2)
for x in ['支持从视频中提取完整音轨。','输出格式支持MP3、WAV、M4A；码率支持128、192、320 kbps。','音频导出按钮和状态使用绿色语义。','不包含AI人声、伴奏、鼓、贝斯等音源分离；若未来实现需另接Demucs类模型。']:
    add_bullet(x)

doc.add_heading('9. 片段过渡', level=1)
doc.add_paragraph('简单过渡属于适合本工具定位的高频增强，但必须保持轻量，不扩展为特效系统。')
for x in ['片段之间可提供“无过渡、淡化、黑场”三个简洁选项。','视频过渡使用FFmpeg xfade；音频衔接使用acrossfade。','常用时长可提供0.2、0.5、1、2秒。','音频裁切点可默认应用20至50毫秒的短交叉淡化，以减少爆音。','过渡会造成相邻片段时间重叠，最终时长计算必须准确。']:
    add_bullet(x)

doc.add_heading('10. 导出设置', level=1)
add_table(['设置项','要求'],[
    ('视频编码','支持H.264，输出MP4；作为推荐默认值'),
    ('分辨率','跟随原视频、2160p、1080p、720p、480p'),
    ('帧率','常见选择：24、25、30、50、60 fps'),
    ('码率/画质','支持码率选择；默认界面可优先显示“节省空间、推荐、高质量”，具体码率放更多设置'),
    ('画幅','原始、16:9、9:16、1:1；裁剪铺满或黑边适配'),
    ('音频','MP3、WAV、M4A；128/192/320 kbps'),
    ('保存位置','“每次导出时选择”或“使用默认文件夹”二选一'),
    ('多时间轴','每条非空时间轴分别输出文件'),
],[1.35,5.15])
for x in ['导出显示真实进度和取消按钮。','导出失败必须显示可理解的FFmpeg错误摘要。','H.264与合理码率、帧率、音频参数应有可直接使用的默认值。','不要让低频技术参数长期占据大量界面空间。']:
    add_bullet(x)

doc.add_heading('11. UI与布局规范', level=1)
doc.add_heading('11.1 总体布局', level=2)
for x in [
    '保持一屏式三段结构：预览与入出点、片段排序、常用输出设置。',
    '预览区是主要视觉中心；音频模式在同一区域切换为详细波形，不引入另一套复杂界面。',
    '右侧保持单面板，按“当前片段”和“导出文件”分为短区块。',
    '顶部中间不显示重复的“添加视频/添加音频”按钮；导入入口放在空白区和时间轴添加卡片。',
    '取消右侧浏览器风格纵向滚动条；面板应通过紧凑排布、折叠低频选项或自定义滚动样式解决。',
]: add_bullet(x)
doc.add_heading('11.2 视觉语言', level=2)
add_table(['颜色','固定语义'],[
    ('蓝紫色','主按钮、播放、选中状态、进度条'),
    ('绿色','入出点有效区间、音频导出、成功状态'),
    ('红色','删除、关闭、错误'),
    ('黄色','代理生成、警告'),
    ('灰色','辅助信息、禁用状态'),
],[1.35,5.15])
for x in ['深色、现代、简洁，延续早期版本视觉风格。','字体比早期默认增大一级，但不能导致面板拥挤。','减少说明文字、重复标签、边框和同时展示的设置。','所有滚动条、进度条、选择控件保持主题化，不出现浏览器默认样式。','导出音频图标使用绿色；播放按钮使用主体蓝紫色。']:
    add_bullet(x)

doc.add_heading('12. 技术实现约束', level=1)
for x in [
    '使用Electron与Web前端实现Windows/macOS桌面端。',
    '使用FFmpeg负责探测、代理、缩略图、裁切、拼接、转码、画幅、音频提取、波形PCM采样和响度处理。',
    '本地媒体通过安全自定义协议访问；不能通过关闭整体Web安全来绕过问题。',
    '缩略图应在主进程中由FFmpeg抓帧，避免Canvas跨源污染导致静默失败。',
    '长音频和特殊格式波形应在主进程中解码为PCM并计算峰值，避免仅依赖浏览器decodeAudioData。',
    '专业相机素材预览使用缓存代理，缓存键应考虑路径、文件大小和修改时间。',
    '最终导出始终读取原始素材，不使用低质量代理作为输出源。',
    '变换、音量、入出点等数据必须存储在具体片段上；分辨率、帧率、编码和保存策略属于导出设置。',
]: add_bullet(x)

doc.add_heading('13. 验收与实际检测标准', level=1)
doc.add_paragraph('“编译通过”或“进程存在”不等于功能验收通过。每次交付必须尽量使用真实媒体完成以下操作链路。')
checks = [
    ('导入','拖入常见MP4/MOV、索尼MXF/MTS以及音频文件；片段正确出现'),
    ('选中预览','逐个点击片段，视频画面或真实音频波形立即显示'),
    ('缩略图','FFmpeg首帧提取成功；失败状态明确且可重试'),
    ('入出点','拖动两端手柄，画面/长光标实时跟随，绿色区间准确'),
    ('播放','仅播放入点至出点，出点自动暂停'),
    ('分割','分割后两段均可独立预览、排序、删除与导出'),
    ('排序','拖动时其他片段平滑让位；导出顺序与UI一致'),
    ('片段变换','缩放/旋转/翻转只影响当前片段，预览与导出一致'),
    ('音频','真实波形、滚轮缩放、中键平移、dB音量、响度统一有效'),
    ('多时间轴','新增、切换、清空/关闭正常；每条非空时间轴分别导出'),
    ('导出','分辨率、帧率、画幅、H.264、音频格式和保存策略实际生效'),
    ('异常处理','无效文件、代理失败、波形失败、导出失败均给出明确反馈'),
]
add_table(['验收项目','通过条件'],checks,[1.35,5.15])
doc.add_heading('13.1 不得宣称“已验证”的情况', level=2)
for x in ['只完成Vite或语法编译。','只确认Electron窗口打开或进程存在。','缩略图、波形仍是占位内容。','按钮存在但没有执行真实媒体操作。','只测试一种示例格式便宣称全部格式可用。','没有真实素材时，将静态检查描述为完整功能验证。']:
    add_bullet(x)

doc.add_heading('14. 开发优先级', level=1)
add_table(['优先级','范围'],[
    ('P0 必须可靠','导入、选中预览、首帧、真实波形、入出点、区间播放、分割、排序、导出'),
    ('P1 高频体验','片段级变换、音量/响度、画幅适配、帧率、保存策略、多时间轴独立导出'),
    ('P2 轻量增强','简单淡化/黑场过渡、画质预设、代理缓存管理、未导出提醒'),
    ('不进入范围','多轨、关键帧、字幕、调色、复杂特效、专业混音、PSD/图片编辑'),
],[1.35,5.15])

doc.add_heading('15. 最终交付原则', level=1)
for x in [
    '严格按照用户明确要求实现，不擅自改变产品定位或增加复杂结构。',
    '功能增加不能牺牲早期版本的简洁与美观。',
    '低级关键链路问题（如选中片段无法预览）必须在交付前发现。',
    '无法使用真实素材验证的功能必须明确标注验证边界，不得夸大完成度。',
    '每次更新应交付可运行Windows版本，并说明本轮实现、验证项目和仍有限制。',
]: add_bullet(x)

doc.core_properties.title = '剪影工坊 - 产品与开发需求规格'
doc.core_properties.subject = '日常视频与音频素材快速处理工具需求汇总'
doc.core_properties.author = '产品需求汇总'
doc.core_properties.keywords = '视频工具, 音频工具, FFmpeg, Electron, 需求规格'
doc.save(OUT)
print(OUT)
